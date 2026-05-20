from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Speaker3_Script.docx")


SECTIONS = [
    {
        "slide": "Slide 2",
        "title": "Presentation Flow Aligned with Team Responsibilities",
        "english": (
            "Before moving into my dashboard section, I want to anchor where it sits in the full presentation. "
            "The story starts with the project goal and data cleaning, then moves into the database design. "
            "My part comes next: I show how the cleaned data becomes an interactive dashboard. After that, "
            "the team continues with statistical testing, regression and prediction, then closes with insights, "
            "export and the live demo."
        ),
        "chinese": (
            "这一页先说明整体汇报路线。我的部分位于数据库之后、统计和预测之前，作用是把已经清洗好的数据"
            "转化成可以交互探索的 dashboard。"
        ),
    },
    {
        "slide": "Slide 8",
        "title": "System Interface: One Entry Point, Eight Analytical Views",
        "english": (
            "This is the main dashboard interface. The important design idea is shared state. The filter rail on "
            "the left controls the year range, genres and minimum votes. Once those controls change, the KPI row, "
            "all charts and the export result update together. So users do not need to run separate queries; they "
            "can explore the same movie slice across Overview, Revenue, Genres and Time."
        ),
        "chinese": (
            "这一页讲系统整体界面。重点是左侧筛选条件会同步影响 KPI、图表和导出结果，说明 dashboard 是一个统一的交互入口。"
        ),
    },
    {
        "slide": "Slide 9",
        "title": "Revenue View: Budget Lifts Revenue, but ROI Remains a Risk Question",
        "english": (
            "In the Revenue view, I first use the budget-versus-revenue scatter plot to explain market scale. "
            "The live view shows a positive relationship: bigger budgets generally create higher box office potential. "
            "But the second point is just as important. Revenue is not the same as efficiency. ROI can still vary widely, "
            "so the dashboard separates the question 'Can this film become large?' from 'Can this investment return well?'"
        ),
        "chinese": (
            "这一页讲 Revenue 视图。预算和票房正相关，但这只说明规模，不等于投资效率；所以要同时看 ROI 风险。"
        ),
    },
    {
        "slide": "Slide 10",
        "title": "Genres and Time: Category Mix and Release Window Shape the Readout",
        "english": (
            "For Genres and Time, the dashboard moves from one financial relationship to practical comparison. "
            "The genre ranking shows which categories perform better under the current filters. The rating box plots "
            "show distribution, not only an average score. Finally, the monthly revenue curve shows how release timing "
            "changes expected revenue. Together, these views help users compare category strategy and release-window strategy."
        ),
        "chinese": (
            "这一页把类型和档期结合起来讲。类型图用于比较 ROI 或平均票房，评分箱线图看口碑分布，月份曲线看上映时机。"
        ),
    },
]


def set_east_asia(run, font="Microsoft YaHei") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_style_font(style, name, size=None, color=None) -> None:
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor(*color)
    if style._element.rPr is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border_cell(cell, color="DADCE0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        element = tc_pr.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_pr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_cell(cell, font_size=8.6, bold=False) -> None:
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(2)
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = bold
            set_east_asia(run)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    set_style_font(doc.styles["Normal"], "Arial", 10.5, (0, 0, 0))
    set_style_font(doc.styles["Title"], "Arial", 22, (0, 0, 0))
    set_style_font(doc.styles["Heading 1"], "Arial", 14, (31, 77, 120))

    p = doc.add_paragraph()
    p.style = "Title"
    run = p.add_run("Speaker 3 Presentation Script")
    set_east_asia(run)

    p = doc.add_paragraph()
    run = p.add_run("ZHOU Can · Slides 2, 8-10 · Dashboard Walkthrough")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Timing target: approximately 3 minutes for the Speaker 3 section, plus a short Talk Map transition if needed.")
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    for cell, text in zip(table.rows[0].cells, ["Slide", "Title", "Speaking focus"]):
        cell.text = text
        shade_cell(cell, "E8EEF5")
        border_cell(cell)
        format_cell(cell, 9, True)

    for item in SECTIONS:
        cells = table.add_row().cells
        values = [item["slide"].replace("Slide ", ""), item["title"], item["chinese"]]
        for cell, text in zip(cells, values):
            cell.text = text
            border_cell(cell)
            format_cell(cell)

    doc.add_paragraph()

    for item in SECTIONS:
        h = doc.add_paragraph()
        h.style = "Heading 1"
        run = h.add_run(f"{item['slide']} - {item['title']}")
        run.bold = True
        set_east_asia(run)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item["english"])
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        set_east_asia(run)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run("中文简释：" + item["chinese"])
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(9.3)
        run.font.color.rgb = RGBColor(85, 85, 85)
        set_east_asia(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "Delivery note: keep the pace conversational. The key contrast is scale versus efficiency in Revenue, "
        "then category and timing strategy in Genres and Time."
    )
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(run)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
