from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Final_Report.docx"
BACKUP = ROOT / "Final_Report.before_final_cleanup.docx"


def set_run_font(paragraph, size=None, bold=None, color=None, name="Aptos"):
    for run in paragraph.runs:
        run.font.name = name
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        if color is not None:
            run.font.color.rgb = RGBColor(*color)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_text(paragraph, text, style=None, size=None, bold=None, color=None):
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def main():
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    doc = Document(REPORT)

    # Cover polish.
    cover_updates = {
        "IMDB Movie Analytics Dashboard": ("IMDB Movie Analytics Dashboard", 22, True, (31, 54, 86)),
        "Final Report": ("Final Report", 16, True, (60, 72, 88)),
        "2026-05-19": ("2026-05-20", 11, False, (80, 80, 80)),
    }
    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        if raw in cover_updates:
            text, size, bold, color = cover_updates[raw]
            set_text(paragraph, text, size=size, bold=bold, color=color)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif raw in {"Computer Science and Technology Program", "United International College"}:
            set_run_font(paragraph, size=10.5, color=(80, 80, 80))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Improve contribution wording while preserving the table.
    if doc.tables:
        table = doc.tables[0]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    set_run_font(p, size=9.5)
        header = table.rows[0]
        for cell in header.cells:
            for p in cell.paragraphs:
                set_run_font(p, size=9.5, bold=True, color=(31, 54, 86))

        replacements = {
            "ER Diagram, Regression and prediction module.": "ER diagram support, regression analysis, and prediction module implementation.",
            "Data preparation and cleaned dataset organization.": "Data preparation, cleaning scripts, cleaned dataset organization, and source-data consistency checks.",
            "Dashboard UI integration, insights, export, and reporting support.": "Dashboard UI integration, insights view, export workflow, and report/presentation support.",
            "Chi-square statistical testing module.": "Chi-square statistical testing module, categorical variable construction, and result interpretation.",
            "ER implementation, MySQL import and validation, SQL export.": "ER implementation, MySQL import and validation, SQL export, and database reproducibility support.",
        }
        for row in table.rows[1:]:
            text = row.cells[1].text.strip()
            if text in replacements:
                row.cells[1].text = replacements[text]
                for p in row.cells[1].paragraphs:
                    set_run_font(p, size=9.5)

    # Replace generated/stale TOC paragraphs with a static final TOC that does not
    # list itself. This avoids broken fields when uploaded to LMS.
    toc_start = toc_end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Table of Contents":
            toc_start = i
        if toc_start is not None and p.text.strip() == "Project Overview":
            toc_end = i
            break
    if toc_start is not None and toc_end is not None:
        toc_paras = doc.paragraphs[toc_start + 1:toc_end]
        static_toc = [
            "1. Project Overview\t3",
            "2. Data Set\t4",
            "3. Features and Approaches\t5",
            "   3.1 Interactive Dashboard and Data Pipeline\t5",
            "   3.2 Statistical Testing, Prediction, Insights and Export\t5",
            "4. Key Findings\t6",
            "5. Conclusion and Discussion\t7",
        ]
        for p, text in zip(toc_paras, static_toc):
            set_text(p, text, style="Normal", size=10.5)
            p.paragraph_format.space_after = Pt(2)
        for p in toc_paras[len(static_toc):]:
            delete_paragraph(p)

    # Add a reproducibility note after the data-generation paragraph if it is not
    # already present.
    marker = "For the Streamlit application, SQL dump files are parsed by Python"
    note = (
        "For LMS submission, the original SQL export is intentionally excluded from the archive to keep the package small. "
        "The cleaned analysis dataset is included, and the data cleaning and dataset-building code is kept in scripts/ and python清洗脚本/ so the process remains reproducible when the original data is available."
    )
    if not any(note in p.text for p in doc.paragraphs):
        for p in doc.paragraphs:
            if p.text.startswith(marker):
                new_p = p.insert_paragraph_before(note)
                p._element.addprevious(new_p._element)
                new_p.style = "Normal"
                set_run_font(new_p, size=10.5)
                break

    # Consistent body rhythm.
    for p in doc.paragraphs:
        if p.style.name == "Normal" and p.text.strip():
            p.paragraph_format.space_after = Pt(6)
            set_run_font(p, size=10.5)
        if "Heading" in p.style.name:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            set_run_font(p, bold=True, color=(31, 54, 86))

    doc.core_properties.title = "IMDB Movie Analytics Dashboard Final Report"
    doc.core_properties.subject = "Database programming final project report"
    doc.core_properties.keywords = "TMDB, IMDB, Streamlit, dashboard, database, final report"
    doc.save(REPORT)
    print(f"Finalized {REPORT.name}")


if __name__ == "__main__":
    main()
