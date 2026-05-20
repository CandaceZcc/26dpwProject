from __future__ import annotations

from docx import Document
from docx.shared import Pt


REPORT = "Final_Report.docx"


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def set_text(paragraph, text):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_after = Pt(2)


def main():
    doc = Document(REPORT)
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Table of Contents":
            start = i
        elif start is not None and p.text.strip() == "Project Overview":
            end = i
            break
    if start is None or end is None:
        raise SystemExit("TOC bounds not found")

    static_toc = [
        "1. Project Overview\t3",
        "2. Data Set\t4",
        "3. Features and Approaches\t5",
        "   3.1 Interactive Dashboard and Data Pipeline\t5",
        "   3.2 Statistical Testing, Prediction, Insights and Export\t5",
        "4. Key Findings\t6",
        "5. Conclusion and Discussion\t7",
    ]
    old = doc.paragraphs[start + 1:end]
    for p, text in zip(old, static_toc):
        set_text(p, text)
    for p in old[len(static_toc):]:
        delete_paragraph(p)
    doc.save(REPORT)
    print("Repaired static report TOC")


if __name__ == "__main__":
    main()
