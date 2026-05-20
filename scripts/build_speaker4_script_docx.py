# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Speaker4_Script.docx")


def set_east_asia(run, font="Microsoft YaHei"):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        element = tc_pr.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_pr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


SLIDES = [
    {
        "slide": "Slide 11",
        "title": "Chi-Square Independence Testing Framework",
        "english": (
            "Here I introduce the Chi-Square module. The goal is simple: after the user applies "
            "dashboard filters, we test whether two categorical variables still look independent, "
            "or whether their observed counts differ enough from the expected counts to suggest an "
            "association. In the dashboard, I use the null hypothesis H0: the two variables are "
            "independent, and I treat p < 0.05 as evidence to reject that assumption."
        ),
        "chinese": (
            "这一页先说明卡方模块的目的：在用户筛选数据后，判断两个分类变量是否仍然独立。"
            "重点讲清楚原假设是“独立”，p 值小于 0.05 时才认为存在显著关联。"
        ),
    },
    {
        "slide": "Slide 12",
        "title": "Derived Variables and Test Configuration",
        "english": (
            "Before running the test, the system converts raw movie fields into categories that are "
            "easier to interpret. Profitability becomes revenue greater than budget; budget is "
            "grouped into tiers; release month becomes a season; and rating becomes a rating tier. "
            "This is important because Chi-Square works on counts, not continuous raw values."
        ),
        "chinese": (
            "这一页解释为什么要做派生字段。卡方检验处理的是频数表，所以预算、月份、评分"
            "这些连续或细粒度字段需要先转成可解释的类别。"
        ),
    },
    {
        "slide": "Slide 13",
        "title": "Profitability Hypotheses",
        "english": (
            "For the profitability tests, the default filtered slice does not show significant "
            "association. Genre versus profitability has p = 0.7402, budget tier versus profitability "
            "has p = 0.7703, and release season versus profitability has p = 0.6589. So in this slice, "
            "we fail to reject H0. A cautious interpretation is that profitability differences are not "
            "strong enough to be explained by these categories alone."
        ),
        "chinese": (
            "这一页讲盈利性相关检验。三个 p 值都大于 0.05，所以不能拒绝独立性假设。"
            "表达时要谨慎：不是说完全没有关系，而是当前筛选条件下证据不足。"
        ),
    },
    {
        "slide": "Slide 14",
        "title": "Rating-Tier Hypotheses",
        "english": (
            "The rating-tier tests tell a different story. Genre versus rating tier and budget tier "
            "versus rating tier both have p-values below 0.0001, so the evidence is much stronger here. "
            "The Cramer's V values are around 0.20, which means the association is not huge, but it is "
            "consistent enough to be meaningful for analysis."
        ),
        "chinese": (
            "这一页形成对比：评分等级的检验结果很显著，p 值小于 0.0001。Cramer’s V "
            "大约是 0.20，说明关联不是特别强，但已经有实际分析价值。"
        ),
    },
    {
        "slide": "Slide 15",
        "title": "Residual Diagnostics",
        "english": (
            "After the p-value, I use standardized residuals to explain where the association comes "
            "from. A positive residual means a cell appears more often than expected under independence, "
            "while a negative residual means it appears less often. The heatmap and stacked bars help us "
            "avoid saying only 'significant' and instead explain which category combinations drive the result."
        ),
        "chinese": (
            "这一页解释残差诊断。p 值只能告诉我们是否显著，标准化残差和热力图能告诉我们"
            "是哪一些类别组合贡献了差异。"
        ),
    },
    {
        "slide": "Slide 16",
        "title": "Engineering Implementation and Takeaway",
        "english": (
            "From the engineering side, this module is reusable and filter-aware. The function for "
            "derived columns prepares the categorical fields, the Chi-Square function builds the "
            "contingency table and computes the statistic, and the interpretation function turns the "
            "p-value into dashboard text. My final takeaway is that formal testing separates weak "
            "profitability evidence from stronger rating-distribution evidence, which makes the dashboard "
            "more analytical than a static chart collection."
        ),
        "chinese": (
            "最后一页收束到工程实现和结论：模块是可复用、会随筛选条件更新的。最终观点是，"
            "正式统计检验让 dashboard 不只是展示图表，而是能区分证据强弱。"
        ),
    },
]


def set_style_font(style, name, size=None, color=None):
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor(*color)
    if style._element.rPr is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    set_style_font(doc.styles["Normal"], "Arial", 10.5, (0, 0, 0))
    set_style_font(doc.styles["Title"], "Arial", 22, (0, 0, 0))
    set_style_font(doc.styles["Heading 1"], "Arial", 14, (31, 77, 120))

    p = doc.add_paragraph()
    p.style = "Title"
    run = p.add_run("Speaker 4 Presentation Script")
    set_east_asia(run)

    p = doc.add_paragraph()
    run = p.add_run("ZHOU Can · Slides 11-16 · Chi-Square Independence Tests")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(run)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    for cell, text in zip(table.rows[0].cells, ["Slide", "Title", "Speaking focus"]):
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.bold = True
                set_east_asia(run)

    for item in SLIDES:
        cells = table.add_row().cells
        values = [item["slide"].replace("Slide ", ""), item["title"], item["chinese"]]
        for cell, text in zip(cells, values):
            cell.text = text
            set_cell_border(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
                    set_east_asia(run)

    doc.add_paragraph()
    for item in SLIDES:
        h = doc.add_paragraph()
        h.style = "Heading 1"
        run = h.add_run(f"{item['slide']} - {item['title']}")
        run.bold = True
        set_east_asia(run)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item["english"])
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        set_east_asia(run)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run("中文简释：" + item["chinese"])
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(9.3)
        run.font.color.rgb = RGBColor(85, 85, 85)
        set_east_asia(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Timing note: read at a steady pace, this script is designed for approximately two minutes."
    )
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(run)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
